from docx import Document

def genera_word(data, output_path="app/docs/report.docx"):
    doc = Document()
    doc.add_heading("Reporte Generado", level=1)
    
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.save(output_path)
    return output_path
